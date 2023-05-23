import json
from docx import Document
import pandas
import math
from ftfy import fix_encoding
import csv
from datetime import datetime as dt
today = dt.today()


# reading search terms from csv file
file_params = pandas.read_csv('search_terms.csv')
keywords = []
locations = []
for index, row in file_params.iterrows():
    sana = row['keyword']
    location = row['location']
    keywords.append(sana)
    locations.append(location)

def format_json_2_docx():
    # Open the JSON file for reading
    with open('filtered.json', 'r') as file:
        json_data = json.load(file)

    # Add the desired clause at the beginning
    modified_data = [{'data': json_data}]

    # Open the same JSON file for writing (which overwrites the existing file)
    with open('filtered_with_header.json', 'w') as file:
        json.dump(modified_data, file, indent=4)

    # sparate call to read line count
    with open("filtered.json", 'r', encoding="utf-8") as fp:
        for count, line in enumerate(fp):
            pass
    # print('Total Lines', count + 1)

    number_of_jobs = (count - 2)/17
    jobs = math.floor(number_of_jobs)
    # print(f'Number of jobs = [filtered.json lines/17]: {jobs}')
    location_and_job_number = ""
    subject = "Open Jobs at https://fi.indeed.com - " + today.strftime("%B %d, %Y")
    keyword_list = f"with following keywords: {keywords}"

    # In case locations are missing and default location 'Suomi' is added into list    
    csv_file = "locations_bup.csv"
    data_string = ""
    # Open the CSV file in read mode
    with open(csv_file, mode='r') as file:
        reader = csv.reader(file)
        # Iterate over each row in the CSV file
        for row in reader:
            # Join the values of each row into a string
            row_string = ','.join(row)
            data_string += row_string
    # print("Data read from CSV:\n", data_string)
    # print(data_string)
    # location_and_job_number = f"At locations:  {data_string} {jobs} open jobs"

    # Load JSON data from the file and add a JSON header
    with open('filtered_with_header.json', 'r') as json_file: 
        json_data = json.load(json_file)

    # Create a new Word document
    document = Document()

    # Iterate through JSON data and add content to the document
    for item in json_data:
        # Add the document title
        document.add_heading(subject, level=1)
        document.add_paragraph(location_and_job_number)
        document.add_paragraph(keyword_list)

        # Add sections to the document
        data = item['data']
        for section in data:
            company_name = section['company_name']
            job_title = section['job_title']
            date = section['date']
            job_location = section['job_location']
            job_url = section['job_url']

            # Add sections to docx
            document.add_heading(company_name, level=1)
            document.add_paragraph(job_title)        
            document.add_paragraph(job_url)
            document.add_paragraph(date)
            document.add_paragraph(job_location)

    document.save('jobs2.docx')
